"""
finetune/parser.py
==================
Multi-format file parser. Extracts raw medical text from:
  PDF, TXT, CSV, DOCX, Images (OCR via pytesseract / Groq Vision fallback)
"""

import os
import io
import base64
from pathlib import Path
from typing import Optional
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".csv", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}


def parse_file(file_path: str) -> dict:
    """
    Parse any supported file and return extracted text.

    Returns:
        {
          "text": str,
          "source_type": str,
          "page_count": int,
          "char_count": int,
          "file_name": str
        }
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")

    print(f"[Parser] Parsing {path.name} (type={ext})")

    if ext == ".pdf":
        result = _parse_pdf(file_path)
    elif ext == ".txt":
        result = _parse_txt(file_path)
    elif ext == ".csv":
        result = _parse_csv(file_path)
    elif ext == ".docx":
        result = _parse_docx(file_path)
    elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}:
        result = _parse_image(file_path)
    else:
        raise ValueError(f"Unsupported extension: {ext}")

    result["file_name"] = path.name
    result["char_count"] = len(result.get("text", ""))
    print(f"[Parser] ✅ Extracted {result['char_count']:,} characters from {path.name}")
    return result


# ─────────────────────────────────────────────
# Individual parsers
# ─────────────────────────────────────────────

def _parse_pdf(file_path: str) -> dict:
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return {
            "text": "\n\n".join(pages),
            "source_type": "pdf",
            "page_count": len(reader.pages)
        }
    except ImportError:
        raise ImportError("pypdf not installed. Run: pip install pypdf")


def _parse_txt(file_path: str) -> dict:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return {"text": text, "source_type": "txt", "page_count": 1}


def _parse_csv(file_path: str) -> dict:
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding="utf-8", errors="ignore")
        
        # Heuristic: find columns with medical content
        text_cols = []
        priority_keywords = [
            "symptom", "diagnosis", "condition", "complaint", "history",
            "finding", "impression", "description", "note", "text",
            "question", "answer", "instruction", "output", "input",
            "clinical", "patient", "disease", "treatment", "medication"
        ]
        
        for col in df.columns:
            col_lower = col.lower()
            if any(kw in col_lower for kw in priority_keywords):
                text_cols.append(col)
        
        # Fallback: use all string columns
        if not text_cols:
            text_cols = [c for c in df.columns if df[c].dtype == "object"]
        
        # If we have symptom + diagnosis columns, format as pairs
        symptom_col = next((c for c in text_cols if "symptom" in c.lower() or "complaint" in c.lower() or "question" in c.lower() or "instruction" in c.lower()), None)
        diagnosis_col = next((c for c in text_cols if "diagnosis" in c.lower() or "condition" in c.lower() or "answer" in c.lower() or "output" in c.lower()), None)
        
        lines = []
        if symptom_col and diagnosis_col:
            for _, row in df.iterrows():
                s = str(row.get(symptom_col, "")).strip()
                d = str(row.get(diagnosis_col, "")).strip()
                if s and d and s != "nan" and d != "nan":
                    lines.append(f"Symptoms: {s}\nDiagnosis: {d}")
        else:
            # Concatenate all content
            for _, row in df.iterrows():
                parts = []
                for col in text_cols:
                    val = str(row.get(col, "")).strip()
                    if val and val != "nan":
                        parts.append(f"{col}: {val}")
                if parts:
                    lines.append(" | ".join(parts))
        
        text = "\n\n".join(lines)
        return {"text": text, "source_type": "csv", "page_count": 1, "row_count": len(df)}
    except ImportError:
        raise ImportError("pandas not installed. Run: pip install pandas")


def _parse_docx(file_path: str) -> dict:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return {
            "text": "\n\n".join(paragraphs),
            "source_type": "docx",
            "page_count": 1
        }
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")


def _parse_image(file_path: str) -> dict:
    """Try pytesseract OCR first, then Groq Vision fallback."""
    text = ""
    method = "none"
    
    # Try pytesseract
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        if len(text.strip()) > 50:
            method = "ocr"
    except Exception as e:
        print(f"[Parser] OCR failed ({e}), trying Groq Vision...")
    
    # Fallback: Groq Vision
    if len(text.strip()) < 50:
        try:
            groq_api_key = os.getenv("GROQ_API_KEY")
            if not groq_api_key:
                raise ValueError("GROQ_API_KEY not set")
            
            client = Groq(api_key=groq_api_key)
            with open(file_path, "rb") as f:
                img_data = base64.b64encode(f.read()).decode()
            
            ext = Path(file_path).suffix.lower().lstrip(".")
            mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
            
            response = client.chat.completions.create(
                model="meta-llama/llama-4-maverick-17b-128e-instruct",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_data}"}},
                        {"type": "text", "text": "Extract all medical text from this image. Include symptoms, diagnoses, lab values, findings, and any clinical notes. Output only the extracted text."}
                    ]
                }],
                max_tokens=2000
            )
            text = response.choices[0].message.content
            method = "groq_vision"
        except Exception as e:
            print(f"[Parser] Groq Vision failed: {e}")
            text = ""
    
    return {
        "text": text,
        "source_type": "image",
        "page_count": 1,
        "ocr_method": method
    }


# ─────────────────────────────────────────────
# CLI test mode
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if "--test" in sys.argv:
        # Create a sample txt file to test
        sample_path = "/tmp/test_medical.txt"
        with open(sample_path, "w") as f:
            f.write("""Patient: 45 year old male
Symptoms: chest pain, shortness of breath, sweating
Duration: 2 hours
History: hypertension, diabetes
Diagnosis: Acute Myocardial Infarction
Treatment: Aspirin, Clopidogrel, PCI

Patient: 28 year old female
Symptoms: severe headache, photophobia, neck stiffness, fever
Duration: 1 day
History: none
Diagnosis: Bacterial Meningitis
Treatment: IV Ceftriaxone, Dexamethasone
""")
        result = parse_file(sample_path)
        print(f"\n✅ Parser test passed!")
        print(f"   Type: {result['source_type']}")
        print(f"   Chars: {result['char_count']}")
        print(f"   Preview: {result['text'][:100]}...")
